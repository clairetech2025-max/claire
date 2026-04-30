#!/usr/bin/env python3
"""
Claire Parser v1
----------------
Multi-ingest parser for the Claire stack.

What it does:
- Recursively scans files and folders
- Safely expands ZIP files into a temp workspace
- Extracts text from:
    - txt, md, json, csv, log, py, html, xml, yaml, yml
    - pdf
    - docx
    - odt
- OCR fallback with Tesseract for:
    - images
    - image-only PDFs
- Transcribes:
    - mp3, wav, m4a, flac, ogg
    - mp4, mov, mkv, avi, webm (audio extraction first)
- Chunks extracted text into memory-safe units
- Writes JSONL capsules with provenance metadata

Design rules:
- The LLM is not the warehouse.
- The parser creates usable memory objects first.
- Provenance is preserved on every chunk.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import logging
import mimetypes
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import zipfile
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Dict, Iterable, Iterator, List, Optional, Tuple

# ---------- optional deps ----------
try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    from odf import teletype
    from odf.opendocument import load as odf_load
except Exception:
    odf_load = None
    teletype = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    from PIL import Image
except Exception:
    Image = None

try:
    import pytesseract
except Exception:
    pytesseract = None

# faster-whisper is optional and preferred if installed
try:
    from faster_whisper import WhisperModel
except Exception:
    WhisperModel = None

# fallback whisper module if present
try:
    import whisper
except Exception:
    whisper = None


SUPPORTED_TEXT = {
    ".txt",
    ".md",
    ".json",
    ".csv",
    ".log",
    ".py",
    ".html",
    ".htm",
    ".xml",
    ".yaml",
    ".yml",
    ".ini",
    ".cfg",
    ".toml",
    ".rtf",
}
SUPPORTED_DOCS = {".pdf", ".docx", ".odt"}
SUPPORTED_IMAGES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"}
SUPPORTED_AUDIO = {".mp3", ".wav", ".m4a", ".flac", ".ogg"}
SUPPORTED_VIDEO = {".mp4", ".mov", ".mkv", ".avi", ".webm"}
SUPPORTED_ARCHIVES = {".zip"}

DEFAULT_EXCLUDE_DIRS = {
    ".git",
    "__pycache__",
    ".venv",
    "venv",
    "node_modules",
    ".cache",
    ".mypy_cache",
    ".pytest_cache",
}


@dataclass
class ChunkRecord:
    chunk_id: str
    title: str
    source_path: str
    parent_archive: Optional[str]
    source_type: str
    doc_type: str
    extraction_method: str
    sequence: int
    total_chunks: int
    text: str
    sha256: str
    file_sha256: str
    created_at_unix: float
    modified_at_unix: float
    size_bytes: int
    page_number: Optional[int] = None
    media_seconds: Optional[float] = None
    language: Optional[str] = None
    notes: Optional[str] = None


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def safe_slug(text: str) -> str:
    text = re.sub(r"[^\w\-\.]+", "_", text.strip())
    return text[:120] if text else "untitled"


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(text: str, max_words: int = 220, overlap_words: int = 40) -> List[str]:
    words = text.split()
    if not words:
        return []

    chunks: List[str] = []
    step = max(1, max_words - overlap_words)
    for start in range(0, len(words), step):
        block = words[start : start + max_words]
        if not block:
            break
        chunks.append(" ".join(block))
        if start + max_words >= len(words):
            break
    return chunks


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def read_text_file(path: Path) -> str:
    if path.suffix.lower() == ".csv":
        rows = []
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(" | ".join(row))
        return "\n".join(rows)

    if path.suffix.lower() == ".json":
        with path.open("r", encoding="utf-8", errors="ignore") as f:
            obj = json.load(f)
        return json.dumps(obj, indent=2, ensure_ascii=False)

    with path.open("r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_docx_text(path: Path) -> str:
    if docx is None:
        raise RuntimeError("python-docx is not installed")
    d = docx.Document(str(path))
    parts: List[str] = []
    for p in d.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)

    for table in d.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


def extract_odt_text(path: Path) -> str:
    if odf_load is None or teletype is None:
        raise RuntimeError("odfpy is not installed")

    doc = odf_load(str(path))
    parts: List[str] = []

    if hasattr(doc, "text") and doc.text is not None:
        try:
            txt = teletype.extractText(doc.text)
            if txt and txt.strip():
                parts.append(txt)
        except Exception:
            pass

    if hasattr(doc, "spreadsheet") and doc.spreadsheet is not None:
        try:
            txt = teletype.extractText(doc.spreadsheet)
            if txt and txt.strip():
                parts.append(txt)
        except Exception:
            pass

    if hasattr(doc, "presentation") and doc.presentation is not None:
        try:
            txt = teletype.extractText(doc.presentation)
            if txt and txt.strip():
                parts.append(txt)
        except Exception:
            pass

    result = "\n\n".join(p for p in parts if p.strip())
    return result.strip()


def extract_pdf_text_native(path: Path) -> Tuple[str, bool]:
    """
    Returns (text, has_text)
    """
    if PyPDF2 is None:
        raise RuntimeError("PyPDF2 is not installed")

    parts: List[str] = []
    has_text = False
    with path.open("rb") as f:
        reader = PyPDF2.PdfReader(f)
        for i, page in enumerate(reader.pages, start=1):
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            txt = clean_text(txt)
            if txt:
                has_text = True
                parts.append(f"[PAGE {i}]\n{txt}")
    return "\n\n".join(parts), has_text


def ocr_image(path: Path) -> str:
    if pytesseract is None or Image is None:
        raise RuntimeError("pytesseract and pillow are required for OCR")
    with Image.open(path) as img:
        return pytesseract.image_to_string(img)


def pdf_to_images(path: Path, out_dir: Path) -> List[Path]:
    """
    Uses pdftoppm if available.
    """
    ensure_dir(out_dir)
    prefix = out_dir / "page"
    cmd = ["pdftoppm", "-png", str(path), str(prefix)]
    subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    return sorted(out_dir.glob("page-*.png"))


def extract_pdf_with_ocr(path: Path, temp_dir: Path) -> str:
    img_dir = temp_dir / f"{safe_slug(path.stem)}_pdfimgs"
    pages = pdf_to_images(path, img_dir)
    parts: List[str] = []
    for i, img_path in enumerate(pages, start=1):
        txt = clean_text(ocr_image(img_path))
        if txt:
            parts.append(f"[PAGE {i} OCR]\n{txt}")
    return "\n\n".join(parts)


def ffprobe_duration(path: Path) -> Optional[float]:
    cmd = [
        "ffprobe",
        "-v",
        "error",
        "-show_entries",
        "format=duration",
        "-of",
        "default=noprint_wrappers=1:nokey=1",
        str(path),
    ]
    try:
        out = subprocess.run(cmd, check=True, capture_output=True, text=True)
        return float(out.stdout.strip())
    except Exception:
        return None


class WhisperTranscriber:
    def __init__(
        self, model_name: str = "base", device: str = "cpu", compute_type: str = "int8"
    ) -> None:
        self.model_name = model_name
        self.device = device
        self.compute_type = compute_type
        self.backend = None
        self.model = None

        if WhisperModel is not None:
            self.backend = "faster-whisper"
            self.model = WhisperModel(
                model_name, device=device, compute_type=compute_type
            )
        elif whisper is not None:
            self.backend = "whisper"
            self.model = whisper.load_model(model_name)
        else:
            raise RuntimeError(
                "No whisper backend installed. Install faster-whisper or openai-whisper."
            )

    def transcribe(self, path: Path) -> Tuple[str, Optional[str]]:
        if self.backend == "faster-whisper":
            segments, info = self.model.transcribe(str(path), vad_filter=True)
            text = " ".join(seg.text.strip() for seg in segments if seg.text.strip())
            language = getattr(info, "language", None)
            return clean_text(text), language

        result = self.model.transcribe(str(path))
        return clean_text(result.get("text", "")), result.get("language")


def extract_audio_from_video(video_path: Path, out_audio_path: Path) -> None:
    cmd = [
        "ffmpeg",
        "-y",
        "-i",
        str(video_path),
        "-vn",
        "-acodec",
        "pcm_s16le",
        "-ar",
        "16000",
        "-ac",
        "1",
        str(out_audio_path),
    ]
    subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)


def guess_doc_type(path: Path) -> str:
    name = path.name.lower()
    if "manifest" in name or "origin" in name:
        return "origin"
    if "parser" in name or "crown jewel" in name or "palantir" in name:
        return "parser"
    if "are" in name:
        return "memory"
    if "diode" in name or "echo" in name or "gov" in name:
        return "governance"
    if "lycan" in name or "spec ops" in name or "chameleon" in name:
        return "defense"
    if "codemask" in name or "ecosystem" in name:
        return "commercial"
    return "general"


def is_supported(path: Path) -> bool:
    s = path.suffix.lower()
    return (
        s in SUPPORTED_TEXT
        or s in SUPPORTED_DOCS
        or s in SUPPORTED_IMAGES
        or s in SUPPORTED_AUDIO
        or s in SUPPORTED_VIDEO
        or s in SUPPORTED_ARCHIVES
    )


class ClaireParser:
    def __init__(
        self,
        output_jsonl: Path,
        temp_root: Path,
        chunk_words: int = 220,
        overlap_words: int = 40,
        whisper_model: str = "base",
        whisper_device: str = "cpu",
        enable_ocr: bool = True,
        enable_media: bool = True,
        logger: Optional[logging.Logger] = None,
    ) -> None:
        self.output_jsonl = output_jsonl
        self.temp_root = temp_root
        self.chunk_words = chunk_words
        self.overlap_words = overlap_words
        self.enable_ocr = enable_ocr
        self.enable_media = enable_media
        self.logger = logger or logging.getLogger("claire_parser")

        self.transcriber: Optional[WhisperTranscriber] = None
        if self.enable_media:
            try:
                self.transcriber = WhisperTranscriber(
                    model_name=whisper_model, device=whisper_device
                )
                self.logger.info("Transcriber ready using whisper backend.")
            except Exception as e:
                self.logger.warning("Media transcription disabled: %s", e)
                self.transcriber = None

        ensure_dir(self.output_jsonl.parent)
        ensure_dir(self.temp_root)

    def log(self, msg: str) -> None:
        self.logger.info(msg)

    def write_records(self, records: List[ChunkRecord]) -> None:
        if not records:
            return
        with self.output_jsonl.open("a", encoding="utf-8") as f:
            for rec in records:
                f.write(json.dumps(asdict(rec), ensure_ascii=False) + "\n")

    def build_records(
        self,
        path: Path,
        text: str,
        extraction_method: str,
        parent_archive: Optional[str] = None,
        media_seconds: Optional[float] = None,
        language: Optional[str] = None,
        notes: Optional[str] = None,
    ) -> List[ChunkRecord]:
        text = clean_text(text)
        if not text:
            return []

        chunks = chunk_text(
            text, max_words=self.chunk_words, overlap_words=self.overlap_words
        )
        stat = path.stat()
        file_hash = sha256_file(path)
        title = path.stem
        doc_type = guess_doc_type(path)
        source_type = path.suffix.lower().lstrip(".")

        total = len(chunks)
        records: List[ChunkRecord] = []
        for i, chunk in enumerate(chunks, start=1):
            chunk_hash = sha256_text(chunk)
            rec = ChunkRecord(
                chunk_id=f"{safe_slug(path.stem)}_{i:04d}",
                title=title,
                source_path=str(path),
                parent_archive=parent_archive,
                source_type=source_type,
                doc_type=doc_type,
                extraction_method=extraction_method,
                sequence=i,
                total_chunks=total,
                text=chunk,
                sha256=chunk_hash,
                file_sha256=file_hash,
                created_at_unix=stat.st_ctime,
                modified_at_unix=stat.st_mtime,
                size_bytes=stat.st_size,
                page_number=None,
                media_seconds=media_seconds,
                language=language,
                notes=notes,
            )
            records.append(rec)
        return records

    def parse_path(self, path: Path, parent_archive: Optional[str] = None) -> int:
        suffix = path.suffix.lower()

        try:
            if suffix in SUPPORTED_ARCHIVES:
                return self.parse_zip(path)

            if suffix in SUPPORTED_TEXT:
                text = read_text_file(path)
                self.write_records(
                    self.build_records(
                        path, text, "native_text", parent_archive=parent_archive
                    )
                )
                return 1

            if suffix == ".docx":
                text = extract_docx_text(path)
                self.write_records(
                    self.build_records(
                        path, text, "docx_extract", parent_archive=parent_archive
                    )
                )
                return 1

            if suffix == ".odt":
                text = extract_odt_text(path)
                self.write_records(
                    self.build_records(
                        path, text, "odt_extract", parent_archive=parent_archive
                    )
                )
                return 1

            if suffix == ".pdf":
                text, has_text = extract_pdf_text_native(path)
                if has_text:
                    self.write_records(
                        self.build_records(
                            path,
                            text,
                            "pdf_text_extract",
                            parent_archive=parent_archive,
                        )
                    )
                    return 1
                if self.enable_ocr:
                    ocr_text = extract_pdf_with_ocr(path, self.temp_root)
                    self.write_records(
                        self.build_records(
                            path,
                            ocr_text,
                            "pdf_ocr_tesseract",
                            parent_archive=parent_archive,
                        )
                    )
                    return 1
                self.log(f"Skipped image-only PDF without OCR: {path}")
                return 0

            if suffix in SUPPORTED_IMAGES:
                if not self.enable_ocr:
                    self.log(f"Skipped image OCR disabled: {path}")
                    return 0
                text = ocr_image(path)
                self.write_records(
                    self.build_records(
                        path, text, "image_ocr_tesseract", parent_archive=parent_archive
                    )
                )
                return 1

            if suffix in SUPPORTED_AUDIO:
                if self.transcriber is None:
                    self.log(f"Skipped audio, no transcriber: {path}")
                    return 0
                duration = ffprobe_duration(path)
                text, language = self.transcriber.transcribe(path)
                self.write_records(
                    self.build_records(
                        path,
                        text,
                        "audio_transcription",
                        parent_archive=parent_archive,
                        media_seconds=duration,
                        language=language,
                    )
                )
                return 1

            if suffix in SUPPORTED_VIDEO:
                if self.transcriber is None:
                    self.log(f"Skipped video, no transcriber: {path}")
                    return 0
                wav_path = self.temp_root / f"{safe_slug(path.stem)}.wav"
                extract_audio_from_video(path, wav_path)
                duration = ffprobe_duration(path)
                text, language = self.transcriber.transcribe(wav_path)
                self.write_records(
                    self.build_records(
                        path,
                        text,
                        "video_audio_transcription",
                        parent_archive=parent_archive,
                        media_seconds=duration,
                        language=language,
                    )
                )
                return 1

            self.log(f"Unsupported file type: {path}")
            return 0

        except Exception as e:
            self.logger.exception("Failed parsing %s: %s", path, e)
            return 0

    def parse_zip(self, zip_path: Path) -> int:
        extract_dir = (
            self.temp_root / f"zip_{safe_slug(zip_path.stem)}_{int(time.time())}"
        )
        ensure_dir(extract_dir)
        count = 0

        with zipfile.ZipFile(zip_path, "r") as zf:
            zf.extractall(extract_dir)

        for item in sorted(extract_dir.rglob("*")):
            if item.is_file() and is_supported(item):
                count += self.parse_path(item, parent_archive=str(zip_path))
        return count

    def parse_tree(self, root: Path) -> int:
        total = 0
        if root.is_file():
            return self.parse_path(root)

        for path in sorted(root.rglob("*")):
            if path.is_dir():
                if path.name in DEFAULT_EXCLUDE_DIRS:
                    continue
                continue
            if any(part in DEFAULT_EXCLUDE_DIRS for part in path.parts):
                continue
            if is_supported(path):
                total += self.parse_path(path)
        return total


def build_arg_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="Claire multi-ingest parser")
    p.add_argument("input_path", help="File or directory to parse")
    p.add_argument(
        "--output-jsonl",
        default="/home/LuciusPrime/claire/data/parser_output.jsonl",
        help="Where JSONL chunk records will be written",
    )
    p.add_argument(
        "--temp-root",
        default="/home/LuciusPrime/claire/data/parser_temp",
        help="Temporary extraction/transcoding workspace",
    )
    p.add_argument("--chunk-words", type=int, default=220, help="Words per chunk")
    p.add_argument(
        "--overlap-words", type=int, default=40, help="Overlap words per chunk"
    )
    p.add_argument("--whisper-model", default="base", help="Whisper model name")
    p.add_argument("--whisper-device", default="cpu", help="cpu or cuda")
    p.add_argument("--disable-ocr", action="store_true", help="Disable Tesseract OCR")
    p.add_argument(
        "--disable-media", action="store_true", help="Disable audio/video transcription"
    )
    p.add_argument(
        "--clear-output",
        action="store_true",
        help="Delete existing output JSONL before writing new records",
    )
    return p


def configure_logging() -> None:
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s | %(levelname)s | %(message)s",
        stream=sys.stdout,
    )


def main() -> int:
    configure_logging()
    args = build_arg_parser().parse_args()

    input_path = Path(args.input_path).expanduser().resolve()
    output_jsonl = Path(args.output_jsonl).expanduser().resolve()
    temp_root = Path(args.temp_root).expanduser().resolve()

    if not input_path.exists():
        logging.error("Input path does not exist: %s", input_path)
        return 1

    if args.clear_output and output_jsonl.exists():
        output_jsonl.unlink()

    parser = ClaireParser(
        output_jsonl=output_jsonl,
        temp_root=temp_root,
        chunk_words=args.chunk_words,
        overlap_words=args.overlap_words,
        whisper_model=args.whisper_model,
        whisper_device=args.whisper_device,
        enable_ocr=not args.disable_ocr,
        enable_media=not args.disable_media,
    )

    logging.info("Starting Claire parser")
    logging.info("Input: %s", input_path)
    logging.info("Output JSONL: %s", output_jsonl)

    total = parser.parse_tree(input_path)
    logging.info("Finished. Parsed units: %s", total)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
